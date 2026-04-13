#!/usr/bin/env python3
"""
Genera DOCX de entregables academicos para el curso 1INF54 (PUCP 2026-1).
Proyecto PDDS - Tasf.B2B

Salida (en documentacion/):
  - 01.definicion.prototipo.v03.docx
  - 12.ana.casos.uso.v01.docx
  - 13.ana.modelo.clases.v01.docx
  - 14.ana.secuencias.escenarios.v01.docx
  - 15.ana.reglas.negocio.v01.docx
  - 16.ana.trazabilidad.lee.v01.docx
  - 21.dis.selec.algoritmos.v02.docx

Requiere: pip install python-docx pillow
           java (para PlantUML)
           tools/plantuml.jar
"""

from __future__ import annotations

import subprocess
import tempfile
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# -- paths --
ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "documentacion"
PLANTUML_JAR = Path(__file__).resolve().parent / "plantuml.jar"
BACKEND_ALG = ROOT / "backend" / "src" / "main" / "java" / "com" / "tasfb2b" / "service" / "algorithm"
BACKEND_SVC = ROOT / "backend" / "src" / "main" / "java" / "com" / "tasfb2b" / "service"

# -- constants --
DATE_STR = datetime.now().strftime("%Y-%m-%d")
SEMESTER = "2026-1"
TEAM = [
    ("A1", "Alonso Reyes Samaniego"),
    ("A2", "Rodrigo Holguin Huari"),
    ("A3", "Sebastian Badajoz Garcia Godos"),
    ("A4", "Adrian Picoy Cotrina"),
]
COURSE = "1INF54 - Proyecto de Desarrollo de Software"
PROJECT = "PDDS - Plataforma de Distribucion y Deteccion de Saturacion"


# ============================================================================
#  Helpers
# ============================================================================

def set_cell_shading(cell, color: str):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)


def new_doc() -> Document:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15
    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    return doc


def add_cover(doc: Document, title: str, subtitle: str):
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(COURSE)
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(PROJECT)
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(subtitle)
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_paragraph()
    doc.add_paragraph()

    for code, name in TEAM:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{code}: ")
        run.bold = True
        p.add_run(name)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Semestre {SEMESTER}  |  {DATE_STR}")

    doc.add_page_break()


def add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(table)
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for par in hdr[i].paragraphs:
            for run in par.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(hdr[i], "1F3A5F")
    for r_idx, row_data in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row_data):
            cells[c_idx].text = str(val)
            for par in cells[c_idx].paragraphs:
                for run in par.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()


def add_code(doc: Document, code: str, *, font_size: int = 8):
    for line in code.strip().split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(font_size)
        run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    doc.add_paragraph()


def render_plantuml(puml_code: str) -> Path | None:
    """Render PlantUML code to a PNG image, return path or None on failure."""
    if not PLANTUML_JAR.exists():
        return None
    with tempfile.NamedTemporaryFile(suffix=".puml", delete=False, mode="w", encoding="utf-8") as f:
        f.write(puml_code)
        puml_path = Path(f.name)
    try:
        subprocess.run(
            ["java", "-jar", str(PLANTUML_JAR), "-tpng", "-o", str(puml_path.parent), str(puml_path)],
            capture_output=True, timeout=30,
        )
        png_path = puml_path.with_suffix(".png")
        if png_path.exists():
            return png_path
    except Exception:
        pass
    return None


def add_plantuml(doc: Document, puml_code: str, caption: str = "", width: float = 6.0):
    """Add a PlantUML diagram as an image, or fallback to code block."""
    png = render_plantuml(puml_code)
    if png:
        doc.add_picture(str(png), width=Inches(width))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(caption)
            run.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    else:
        if caption:
            doc.add_paragraph(caption, style="Heading 3")
        add_code(doc, puml_code)
    doc.add_paragraph()


def read_java_file(path: Path) -> str:
    """Read a Java source file and return its content."""
    if path.exists():
        return path.read_text(encoding="utf-8")
    return f"// Archivo no encontrado: {path.name}"


def bullet(doc, text):
    doc.add_paragraph(text, style="List Bullet")


def numbered(doc, text):
    doc.add_paragraph(text, style="List Number")


# ============================================================================
#  01 - Definicion del Prototipo v03
# ============================================================================

def build_doc_01() -> tuple[str, Document]:
    doc = new_doc()
    add_cover(doc, "Definicion del Prototipo", "Tercera Iteracion (v03)")

    doc.add_heading("1. Objetivo del Prototipo", level=1)
    doc.add_paragraph(
        "El prototipo v03 valida las funcionalidades esenciales correspondientes a las "
        "Listas de Exigencias de tipo Exigible (LE tipo E) del proyecto Tasf.B2B. "
        "El sistema cubre: (i) registro de maletas a enviar; (ii) planificacion y "
        "replanificacion de rutas mediante un componente planificador parametrizable; y "
        "(iii) monitoreo grafico de operaciones en mapa interactivo en tiempo real."
    )
    doc.add_paragraph(
        "El componente planificador resuelve mediante parametros los 3 escenarios "
        "evaluados en el curso: operacion dia a dia, simulacion de periodo (semanal, 5 dias "
        "o 3 dias que toma entre 30 y 90 minutos en ejecutarse), y simulacion hasta el "
        "colapso de las operaciones. Presenta de manera grafica informacion relevante del "
        "desempeno en los 3 escenarios usando colores de semaforo (verde, ambar, rojo) "
        "con rangos parametrizables."
    )

    doc.add_heading("2. Alcance Funcional (LE Exigibles)", level=1)
    add_table(doc,
        ["Bloque LE", "Categoria", "Funcionalidad"],
        [
            ["LE-001..LE-028", "Monitoreo", "Dashboard con mapa operativo, visualizacion de envios y vuelos en ruta, KPIs de estado, alertas operacionales, filtros, busqueda y detalle por envio/nodo"],
            ["LE-029..LE-044", "Gestion de envios", "Registro e importacion de envios, planificacion automatica de ruta, estados del envio (PENDING, IN_ROUTE, DELIVERED, CRITICAL, DELAYED), trazabilidad de paradas"],
            ["LE-049..LE-057", "Simulacion", "Configuracion y ejecucion de 3 escenarios, controles de inicio/pausa/detencion/velocidad, generacion de demanda sintetica, exportacion de resultados (CSV/PDF)"],
            ["LE-074", "Reportes", "Cumplimiento SLA por tipo de ruta (intra/intercontinental), por cliente (aerolinea) y por nodo destino (ICAO)"],
        ],
    )

    doc.add_heading("3. Escenarios Implementados", level=1)

    # DAY_TO_DAY
    doc.add_heading("3.1 Operacion Dia a Dia (DAY_TO_DAY)", level=2)
    doc.add_paragraph(
        "Simula el flujo continuo de envios de maletas en tiempo casi real. El motor de "
        "simulacion ejecuta ticks cada 2 segundos con velocidad ajustable (1x-10x). Los "
        "envios llegan continuamente y se planifican en lotes incrementales de hasta 120 "
        "envios por tick. El monitor de colapso ejecuta cada 30 segundos verificando "
        "saturacion del sistema."
    )
    bullet(doc, "Mapa con envios IN_ROUTE y rutas activas con interpolacion de posicion.")
    bullet(doc, "KPIs actualizados cada 2 segundos: envios activos, criticos, entregados, SLA %.")
    bullet(doc, "Semaforo de aeropuertos: verde (NORMAL <70%), ambar (ALERTA 70-90%), rojo (CRITICO >90%).")
    bullet(doc, "Detalle de envio con trazabilidad de paradas y auditoria de eventos.")

    # PERIOD_SIMULATION
    doc.add_heading("3.2 Simulacion de Periodo (PERIOD_SIMULATION)", level=2)
    doc.add_paragraph(
        "Reproduce el comportamiento de varios dias (3, 5 o 7) en una ventana controlada "
        "de 30 a 90 minutos reales, conforme al requisito de la situacion autentica. "
        "La velocidad se calcula automaticamente: velocidad = (dias x 24 x 60) / minutos_ejecucion. "
        "Se genera volumen sintetico con patron NORMAL y se planifica en lotes de hasta 220 envios por tick."
    )
    add_table(doc,
        ["Configuracion", "Dias simulados", "Minutos reales", "Velocidad resultante"],
        [
            ["PERIOD_D3_M30", "3", "30", "x144"],
            ["PERIOD_D5_M60", "5", "60", "x120"],
            ["PERIOD_D7_M90", "7", "90", "x112"],
        ],
    )

    # COLLAPSE_TEST
    doc.add_heading("3.3 Simulacion de Colapso Logistico (COLLAPSE_TEST)", level=2)
    doc.add_paragraph(
        "Mide la capacidad maxima del sistema y el punto de colapso operacional. El patron "
        "de carga escala progresivamente: NORMAL -> PEAK -> COLLAPSE (24+-16 maletas, cada "
        "1-2 minutos, concentracion en hubs americanos). El colapso se declara cuando >90% "
        "de aeropuertos alcanzan estado CRITICO. Planifica en lotes de hasta 320 envios por tick."
    )
    bullet(doc, "Indicador de riesgo de colapso (%) con formula: 60% ocupacion promedio + 30% fraccion nodos criticos + 10% fraccion envios problematicos.")
    bullet(doc, "Tiempo estimado al colapso (horas).")
    bullet(doc, "Lista de nodos cuello de botella ordenados por ocupacion descendente.")

    doc.add_heading("4. Funcionalidades Operativas Transversales", level=1)

    doc.add_heading("4.1 Importacion de datos", level=2)
    doc.add_paragraph(
        "Importacion de aeropuertos, vuelos y envios via CSV. Job asincrono de importacion "
        "masiva del dataset oficial (carpeta datos/envios/) con procesamiento paralelo. "
        "Resultado consistente: filas_leidas = filas_importadas + filas_fallidas."
    )

    doc.add_heading("4.2 Monitoreo en mapa interactivo", level=2)
    doc.add_paragraph(
        "Dashboard con MapLibre: envios con interpolacion de posicion en tiempo real, "
        "rutas completadas (verde) y pendientes (amarillo), estados por color, panel "
        "lateral de KPIs, detalle de envio con paradas, detalle de nodo con ocupacion y vuelos."
    )

    doc.add_heading("4.3 Reportes y exportacion", level=2)
    doc.add_paragraph(
        "Reporte SLA por tipo de ruta, cliente y nodo destino. Exportacion de resultados "
        "operativos en CSV y PDF."
    )

    doc.add_heading("5. Arquitectura Tecnica", level=1)
    add_table(doc,
        ["Capa", "Tecnologia"],
        [
            ["Frontend", "Next.js 16, React 19, TypeScript 5, Tailwind CSS 4, MapLibre GL"],
            ["Backend", "Spring Boot 3.4.4, Java 21, Spring Data JPA"],
            ["Base de datos", "PostgreSQL 15"],
            ["Infraestructura", "Docker Compose"],
        ],
    )

    # Architecture diagram
    add_plantuml(doc, """@startuml
skinparam componentStyle rectangle
skinparam backgroundColor white

package "Frontend (Next.js 16)" {
  [Dashboard / Mapa] as dash
  [Simulacion] as sim
  [Import] as imp
  [Reportes] as rep
}

package "Backend (Spring Boot 3.4.4)" {
  [SimulationController] as sc
  [DashboardController] as dc
  [ShipmentController] as shc
  [ReportsController] as rc
  [SimulationEngineService] as eng
  [RoutePlannerService] as rps
  [GeneticAlgorithm] as ga
  [AntColonyOptimization] as aco
}

database "PostgreSQL 15" as db

dash --> dc : REST
sim --> sc : REST
imp --> shc : REST
rep --> rc : REST
sc --> eng
eng --> rps
rps --> ga
rps --> aco
rps --> db
eng --> db
@enduml""", "Figura 1. Arquitectura general del sistema")

    doc.add_heading("6. Evidencias Minimas", level=1)
    add_table(doc,
        ["Cod.", "Evidencia"],
        [
            ["E01", "Mapa escenario DAY_TO_DAY con envios en ruta"],
            ["E02", "Mapa escenario PERIOD_SIMULATION activo"],
            ["E03", "Mapa escenario COLLAPSE_TEST + indicador de riesgo"],
            ["E04", "Importacion oficial DONE con consistencia de filas"],
            ["E05", "Reporte SLA por ruta/cliente/destino"],
            ["E06", "Export CSV y PDF generados"],
        ],
    )

    doc.add_heading("7. Criterios de Aceptacion", level=1)
    bullet(doc, "Se visualizan los 3 escenarios en el mapa interactivo.")
    bullet(doc, "El componente planificador resuelve mediante parametros los 3 escenarios.")
    bullet(doc, "Se presenta graficamente informacion relevante del desempeno.")
    bullet(doc, "Se usan colores de semaforo con rangos parametrizables.")
    bullet(doc, "Plazos de entrega: 1 dia intracontinental, 2 dias intercontinental.")
    bullet(doc, "La simulacion de periodo toma entre 30 y 90 minutos en ejecutarse.")

    return "01.definicion.prototipo.v03.docx", doc


# ============================================================================
#  12 - Casos de Uso
# ============================================================================

def build_doc_12() -> tuple[str, Document]:
    doc = new_doc()
    add_cover(doc, "Documentacion de Analisis", "12. Casos de Uso (v01)")

    doc.add_heading("1. Actores del Sistema", level=1)
    add_table(doc,
        ["Actor", "Descripcion"],
        [
            ["Operador de simulacion", "Configura escenarios, ejecuta simulacion, monitorea mapa y responde alertas operacionales."],
            ["Analista operacional", "Consulta indicadores de desempeno, verifica cumplimiento SLA y exporta resultados."],
            ["Sistema (Scheduler)", "Motor de simulacion que ejecuta ticks automaticos periodicamente (cada 2 segundos)."],
        ],
    )

    # Use case diagram
    add_plantuml(doc, """@startuml
left to right direction
skinparam backgroundColor white
skinparam actorStyle awesome

actor "Operador" as op
actor "Analista" as an
actor "Sistema\\n(Scheduler)" as sys

rectangle "Tasf.B2B" {
  usecase "CU-01 Configurar\\nescenario" as cu01
  usecase "CU-02 Iniciar/Pausar/\\nDetener simulacion" as cu02
  usecase "CU-03 Monitorear envios\\ny vuelos en mapa" as cu03
  usecase "CU-04 Importar\\nenvios oficiales" as cu04
  usecase "CU-05 Consultar detalle\\nde envio y paradas" as cu05
  usecase "CU-06 Gestionar\\nenvio manual" as cu06
  usecase "CU-07 Consultar\\nreporte SLA" as cu07
  usecase "CU-08 Exportar\\nresultados CSV/PDF" as cu08
}

op --> cu01
op --> cu02
op --> cu03
op --> cu04
op --> cu05
op --> cu06
an --> cu07
an --> cu08
sys --> cu02 : <<trigger>>
@enduml""", "Figura 1. Diagrama de Casos de Uso")

    doc.add_heading("2. Catalogo de Casos de Uso", level=1)
    add_table(doc,
        ["ID", "Caso de Uso", "Actor", "LE"],
        [
            ["CU-01", "Configurar escenario de simulacion", "Operador", "LE-049..LE-057"],
            ["CU-02", "Iniciar/pausar/detener simulacion", "Operador", "LE-049..LE-057"],
            ["CU-03", "Monitorear envios y vuelos en mapa", "Operador", "LE-001..LE-028"],
            ["CU-04", "Importar envios oficiales", "Operador", "LE-029..LE-044"],
            ["CU-05", "Consultar detalle de envio y paradas", "Operador", "LE-029..LE-044"],
            ["CU-06", "Gestionar envio manual", "Operador", "LE-029..LE-044"],
            ["CU-07", "Consultar reporte SLA", "Analista", "LE-074"],
            ["CU-08", "Exportar resultados operativos CSV/PDF", "Analista", "LE-049..LE-057"],
        ],
    )

    # Detailed CU specs
    cases = [
        ("CU-01", "Configurar escenario de simulacion", "Operador", "Sistema disponible. Simulacion detenida (isRunning = false).",
         ["El actor selecciona escenario: DAY_TO_DAY, PERIOD_SIMULATION o COLLAPSE_TEST.",
          "Ajusta parametros: dias, minutos de ejecucion, volumen (media y varianza), multiplicador de vuelos, cancelacion %, umbrales de semaforo, algoritmo primario y secundario.",
          "Solicita guardar configuracion (POST /api/simulation/configure).",
          "Sistema valida que la simulacion no este corriendo y persiste la configuracion."],
         "Si simulacion esta corriendo, sistema rechaza reconfiguraccion.",
         "Configuracion activa registrada y consultable via GET /api/simulation/state."),

        ("CU-02", "Iniciar/Pausar/Detener simulacion", "Operador", "Configuracion valida disponible.",
         ["Actor presiona iniciar (POST /api/simulation/start).",
          "Sistema marca isRunning = true; SimulationRuntimeService registra hora de inicio.",
          "SimulationEngineService comienza ticks periodicos (cada 2 segundos).",
          "Actor puede pausar (POST /pause) o cambiar velocidad (POST /speed).",
          "Actor detiene simulacion (POST /stop); isRunning = false, runtime se reinicia."],
         "Si ya esta corriendo y se intenta iniciar, sistema responde estado actual sin duplicar.",
         "Estado runtime consistente. KPIs actualizados en cada tick."),

        ("CU-03", "Monitorear envios y vuelos en mapa", "Operador", "Simulacion iniciada o datos cargados.",
         ["Actor abre dashboard principal (/).",
          "Sistema muestra mapa MapLibre con capas de rutas, envios y panel KPI.",
          "Mapa actualiza posiciones cada 2 segundos (polling).",
          "Actor filtra por estado (PENDING, IN_ROUTE, DELIVERED, CRITICAL) o busca por codigo.",
          "Actor selecciona envio para ver detalle: codigo, origen, destino, estado, progreso, deadline, paradas.",
          "Actor selecciona nodo para ver ocupacion, vuelos y envios asociados."],
         "Sin datos cargados, mapa muestra red vacia con aeropuertos sin actividad.",
         "Actor identifica envios at-risk, overdue o criticos; aeropuertos en ALERTA o CRITICO."),

        ("CU-04", "Importar envios oficiales", "Operador", "Carpeta de datos oficiales disponible (datos/envios/).",
         ["Actor navega a /import y selecciona tipo de importacion.",
          "Para CSV individual: carga archivo y sistema procesa sincronamente.",
          "Para importacion masiva: POST /api/import/envios-dataset-full; sistema crea job asincrono.",
          "Actor consulta estado del job periodicamente.",
          "Sistema finaliza con resumen: filas_leidas = importadas + fallidas."],
         "Archivo con formato invalido: sistema reporta errores clasificados por causa.",
         "Envios persistidos y listos para planificacion incremental."),

        ("CU-05", "Consultar detalle de envio y paradas", "Operador", "Envios existentes en el sistema.",
         ["Actor busca envio por codigo unico o seleccion en mapa.",
          "Sistema devuelve cabecera: codigo, aerolinea, origen, destino, estado, deadline.",
          "Sistema muestra TravelStops: aeropuerto, vuelo asignado, orden, estado del tramo.",
          "Actor revisa progreso por tramo y tiempos estimados vs reales."],
         "", "Trazabilidad operativa del envio disponible."),

        ("CU-06", "Gestionar envio manual", "Operador", "Aeropuertos y vuelos cargados.",
         ["Actor registra envio (POST /api/shipments) con origen, destino, maletas, aerolinea.",
          "Sistema valida origen != destino y capacidad de almacen.",
          "Sistema crea Shipment con deadline automatico (intra=+1d, inter=+2d) y planifica ruta.",
          "Actor puede confirmar entrega manual (POST /api/simulation/deliver/{id})."],
         "Si planificacion falla, envio queda en estado CRITICAL.",
         "Envio actualizado y evento registrado en ShipmentAuditLog."),

        ("CU-07", "Consultar reporte SLA", "Analista", "Envios DELIVERED existentes.",
         ["Analista navega a /reports.",
          "Selecciona rango de fechas (from/to).",
          "Sistema calcula SLA por: tipo de ruta (intra/inter), cliente (aerolinea, top 6), nodo destino (ICAO, top 8).",
          "Frontend muestra barras de progreso y porcentajes on-time."],
         "Sin datos en el rango, indicadores en cero.",
         "Informe SLA para seguimiento operativo."),

        ("CU-08", "Exportar resultados operativos CSV/PDF", "Analista", "Simulacion ejecutada con resultados.",
         ["Analista solicita exportacion (GET /api/simulation/results/export?format=csv|pdf).",
          "Sistema computa KPIs finales.",
          "Sistema genera archivo con metricas: total envios, entregados, on-time %, transito promedio.",
          "Analista descarga archivo como evidencia."],
         "", "Archivo de evidencia generado y descargado."),
    ]

    doc.add_heading("3. Especificacion de Casos de Uso", level=1)
    for cu_id, cu_name, actor, pre, flow, alt, post in cases:
        doc.add_heading(f"{cu_id} - {cu_name}", level=2)
        p = doc.add_paragraph(); r = p.add_run("Actor principal: "); r.bold = True; p.add_run(actor)
        p = doc.add_paragraph(); r = p.add_run("Precondiciones: "); r.bold = True; p.add_run(pre)
        p = doc.add_paragraph(); r = p.add_run("Flujo principal:"); r.bold = True
        for i, step in enumerate(flow, 1):
            numbered(doc, f"{step}")
        if alt:
            p = doc.add_paragraph(); r = p.add_run("Flujo alterno: "); r.bold = True; p.add_run(alt)
        p = doc.add_paragraph(); r = p.add_run("Postcondiciones: "); r.bold = True; p.add_run(post)
        doc.add_paragraph()

    return "12.ana.casos.uso.v01.docx", doc


# ============================================================================
#  13 - Modelo de Clases
# ============================================================================

def build_doc_13() -> tuple[str, Document]:
    doc = new_doc()
    add_cover(doc, "Documentacion de Analisis", "13. Modelo de Clases del Dominio (v01)")

    doc.add_heading("1. Objetivo", level=1)
    doc.add_paragraph(
        "Definir el modelo de dominio del sistema Tasf.B2B que soporta la simulacion "
        "logistica, planificacion de rutas y monitoreo operacional del traslado de maletas "
        "entre aeropuertos de America, Asia y Europa."
    )

    doc.add_heading("2. Diagrama de Clases", level=1)
    add_plantuml(doc, """@startuml
skinparam backgroundColor white
skinparam classAttributeIconSize 0
skinparam classFontSize 11
hide empty methods

enum Continent {
  AMERICA
  EUROPE
  ASIA
}

enum ShipmentStatus {
  PENDING
  IN_ROUTE
  DELIVERED
  CRITICAL
  DELAYED
}

enum FlightStatus {
  SCHEDULED
  IN_FLIGHT
  COMPLETED
  CANCELLED
}

enum StopStatus {
  PENDING
  IN_TRANSIT
  COMPLETED
}

enum AlgorithmType {
  GENETIC
  ANT_COLONY
  SIMULATED_ANNEALING
}

enum SimulationScenario {
  DAY_TO_DAY
  PERIOD_SIMULATION
  COLLAPSE_TEST
}

class Airport {
  -id: Long
  -icaoCode: String <<unique>>
  -city: String
  -country: String
  -latitude: Double
  -longitude: Double
  -continent: Continent
  -maxStorageCapacity: Integer
  -currentStorageLoad: Integer
  -gmtOffset: Integer
  +getOccupancyPct(): Double
  +getStatus(normal%, warning%): AirportStatus
}

class Flight {
  -id: Long
  -flightCode: String
  -scheduledDeparture: LocalDateTime
  -scheduledArrival: LocalDateTime
  -status: FlightStatus
  -maxCapacity: Integer
  -currentLoad: Integer
  -isInterContinental: Boolean
  +getAvailableCapacity(): Integer
  +getLoadPct(): Double
  +getTransitTimeDays(): Double
}

class Shipment {
  -id: Long
  -shipmentCode: String <<unique>>
  -airlineName: String
  -luggageCount: Integer
  -registrationDate: LocalDateTime
  -deadline: LocalDateTime
  -status: ShipmentStatus
  -deliveredAt: LocalDateTime
  -progressPercentage: Double
  -isInterContinental: Boolean
  +isDeliveredOnTime(): Boolean
}

class TravelStop {
  -id: Long
  -stopOrder: Integer
  -stopStatus: StopStatus
  -scheduledArrival: LocalDateTime
  -actualArrival: LocalDateTime
}

class SimulationConfig {
  -id: Long
  -scenario: SimulationScenario
  -simulationDays: Integer
  -executionMinutes: Integer
  -initialVolumeAvg: Integer
  -initialVolumeVariance: Integer
  -flightFrequencyMultiplier: Double
  -cancellationRatePct: Double
  -normalThresholdPct: Double
  -warningThresholdPct: Double
  -primaryAlgorithm: AlgorithmType
  -secondaryAlgorithm: AlgorithmType
  -isRunning: Boolean
  -startedAt: LocalDateTime
}

class ShipmentAuditLog {
  -id: Long
  -eventType: String
  -message: String
  -eventAt: LocalDateTime
}

class OperationalAlert {
  -id: Long
  -alertType: String
  -message: String
  -status: String
  -createdAt: LocalDateTime
}

Airport "1" <-- "*" Flight : origin
Airport "1" <-- "*" Flight : destination
Airport "1" <-- "*" TravelStop
Shipment "1" *-- "*" TravelStop
Flight "1" <-- "*" TravelStop
Shipment "1" *-- "*" ShipmentAuditLog
Airport "1" -- "*" OperationalAlert
Shipment "*" --> "1" Airport : origin
Shipment "*" --> "1" Airport : destination
Airport -- Continent
Flight -- FlightStatus
Shipment -- ShipmentStatus
TravelStop -- StopStatus
SimulationConfig -- SimulationScenario
SimulationConfig -- AlgorithmType
@enduml""", "Figura 1. Diagrama de Clases del Dominio", width=6.5)

    doc.add_heading("3. Descripcion de Clases", level=1)

    classes_data = [
        ("Airport", "Nodo de la red logistica. Almacen en aeropuerto con capacidad entre 500 y 800 maletas. Estado derivado por ocupacion: NORMAL/ALERTA/CRITICO con umbrales parametrizables."),
        ("Flight", "Vuelo programado entre dos aeropuertos. Capacidad entre 150-250 (intracontinental) y 150-400 (intercontinental). Gestion de carga actual vs maxima."),
        ("Shipment", "Grupo de maletas de una aerolinea en transito. Deadline automatico: +1 dia (intra) o +2 dias (intercontinental). Ciclo de vida: PENDING -> IN_ROUTE -> DELIVERED | CRITICAL | DELAYED."),
        ("TravelStop", "Parada individual en la ruta planificada de un envio. Cada parada asocia un aeropuerto con un vuelo asignado y un orden secuencial."),
        ("SimulationConfig", "Entidad singleton que gobierna el escenario activo. Contiene parametros de simulacion, umbrales de semaforo y algoritmos seleccionados."),
        ("ShipmentAuditLog", "Registro de eventos del ciclo de vida de un envio: creacion, planificacion, partida, llegada, entrega, replanificacion, vencimiento."),
        ("OperationalAlert", "Alertas generadas por el sistema cuando se detectan condiciones criticas en aeropuertos o envios."),
    ]

    for cls_name, cls_desc in classes_data:
        doc.add_heading(f"3.{classes_data.index((cls_name,cls_desc))+1} {cls_name}", level=2)
        doc.add_paragraph(cls_desc)

    doc.add_heading("4. Relaciones Principales", level=1)
    add_table(doc,
        ["Relacion", "Cardinalidad", "Descripcion"],
        [
            ["Airport - Flight", "1..* (origen y destino)", "Un aeropuerto es origen/destino de multiples vuelos"],
            ["Shipment - TravelStop", "1 -> 0..*", "Un envio tiene cero o mas paradas planificadas (composicion)"],
            ["TravelStop - Flight", "* -> 1", "Cada parada usa un vuelo asignado para el tramo"],
            ["TravelStop - Airport", "* -> 1", "Cada parada corresponde a un aeropuerto"],
            ["Shipment - ShipmentAuditLog", "1 -> 0..*", "Eventos de auditoria por envio (composicion)"],
            ["Shipment - Airport", "* -> 1 (origen, destino)", "Aeropuertos de inicio y fin del envio"],
        ],
    )

    return "13.ana.modelo.clases.v01.docx", doc


# ============================================================================
#  14 - Secuencias de Escenarios
# ============================================================================

def build_doc_14() -> tuple[str, Document]:
    doc = new_doc()
    add_cover(doc, "Documentacion de Analisis", "14. Diagramas de Secuencia (v01)")

    doc.add_heading("1. Objetivo", level=1)
    doc.add_paragraph(
        "Documentar las secuencias de interaccion principales del sistema Tasf.B2B "
        "mediante diagramas UML de secuencia."
    )

    # S1: Configure + Start
    doc.add_heading("2. Configurar e Iniciar Simulacion", level=1)
    add_plantuml(doc, """@startuml
skinparam backgroundColor white
actor Operador
participant "Frontend\\n(Simulation)" as FE
participant "SimulationController" as SC
participant "SimulationConfig\\nRepository" as repo
participant "SimulationRuntime\\nService" as SRS

== Configuracion ==
Operador -> FE : Seleccionar escenario\\ny ajustar parametros
FE -> SC : POST /api/simulation/configure
SC -> SC : Validar isRunning == false
SC -> repo : save(config)
SC --> FE : SimulationState (config actualizada)

== Inicio ==
Operador -> FE : Presionar "Iniciar"
FE -> SC : POST /api/simulation/start
SC -> repo : config.setIsRunning(true)
SC -> SRS : markStarted()
SC --> FE : SimulationState (RUNNING)
FE -> FE : Iniciar polling cada 2s
@enduml""", "Figura 1. Secuencia: Configurar e Iniciar Simulacion")

    # S2: Tick
    doc.add_heading("3. Tick de Simulacion (Motor Interno)", level=1)
    doc.add_paragraph(
        "El SimulationEngineService ejecuta un tick cada 2 segundos reales. Cada tick "
        "avanza el reloj simulado y ejecuta 4 fases secuenciales."
    )
    add_plantuml(doc, """@startuml
skinparam backgroundColor white
participant "SimulationEngine\\nService" as ENG
participant "ShipmentRepository" as SR
participant "RoutePlannerService" as RPS
participant "FlightRepository" as FR
participant "TravelStopRepository" as TSR
participant "SimulationRuntime\\nService" as SRS

ENG -> ENG : calcular simulatedNow\\n(horizonte temporal)

== Fase 1: Planificar envios PENDING ==
ENG -> SR : findPendingWithoutRoute\\n(horizon, batchSize)
loop para cada shipment del lote
  ENG -> RPS : planShipment(shipment, algorithm)
  RPS -> RPS : resolveOptimizer()\\n(GA o ACO)
  RPS -> RPS : filtrar vuelos elegibles
  RPS -> RPS : buildDirectCandidate()
  RPS -> RPS : optimizer.planRoute()
  RPS -> RPS : selectBestCandidate()
  RPS -> RPS : allocateFlightLoads()
  RPS -> TSR : saveAll(stops)
  RPS -> FR : saveAll(updatedLoads)
end

== Fase 2: Activar vuelos ==
ENG -> FR : findByStatusAndDeparture\\n(SCHEDULED, horizon)
ENG -> FR : marcar IN_FLIGHT
ENG -> TSR : marcar IN_TRANSIT

== Fase 3: Cerrar vuelos ==
ENG -> FR : findByStatusAndArrival\\n(IN_FLIGHT, horizon)
ENG -> FR : marcar COMPLETED
ENG -> TSR : marcar COMPLETED
ENG -> SR : marcar DELIVERED\\n(si ultima parada)

== Fase 4: Marcar overdue ==
ENG -> SR : markOverdue(horizon)
ENG -> SRS : markTick(horizon)
@enduml""", "Figura 2. Secuencia: Tick de Simulacion", width=6.5)

    # S3: Import
    doc.add_heading("4. Importacion Masiva de Envios", level=1)
    add_plantuml(doc, """@startuml
skinparam backgroundColor white
actor Operador
participant "Frontend\\n(Import)" as FE
participant "DataImport\\nController" as DIC
participant "EnviosImport\\nJobService" as JOB
participant "DataImport\\nService" as DIS
database "PostgreSQL" as DB

Operador -> FE : Iniciar importacion masiva
FE -> DIC : POST /api/import/envios-dataset-full
DIC -> JOB : startFullJob()
JOB -> JOB : crear job RUNNING
DIC --> FE : jobId + estado RUNNING

loop polling cada 3s
  FE -> DIC : GET /api/import/logs
  DIC --> FE : estado del job
end

JOB -> DIS : procesarArchivos(datos/envios/)
loop para cada archivo
  DIS -> DIS : parsear CSV
  DIS -> DB : batch insert shipments
end
JOB -> JOB : estado = DONE\\nleidos = importados + fallidos
FE -> DIC : GET /api/import/logs
DIC --> FE : resumen final
@enduml""", "Figura 3. Secuencia: Importacion Masiva")

    # S4: SLA Report
    doc.add_heading("5. Consulta de Reporte SLA", level=1)
    add_plantuml(doc, """@startuml
skinparam backgroundColor white
actor Analista
participant "Frontend\\n(Reports)" as FE
participant "ReportsController" as RC
participant "ReportingService" as RS
database "PostgreSQL" as DB

Analista -> FE : Seleccionar rango de fechas
FE -> RC : GET /api/reports/sla-compliance\\n?from=...&to=...
RC -> RS : calculateSla(from, to)
RS -> DB : query shipments DELIVERED\\nen rango
RS -> RS : calcular SLA por:\\n- tipo ruta (intra/inter)\\n- cliente (aerolinea)\\n- nodo destino (ICAO)
RS --> RC : SlaComplianceDto
RC --> FE : JSON response
FE -> FE : renderizar barras\\ny porcentajes
@enduml""", "Figura 4. Secuencia: Reporte SLA")

    return "14.ana.secuencias.escenarios.v01.docx", doc


# ============================================================================
#  15 - Reglas de Negocio
# ============================================================================

def build_doc_15() -> tuple[str, Document]:
    doc = new_doc()
    add_cover(doc, "Documentacion de Analisis", "15. Reglas de Negocio (v01)")

    doc.add_heading("1. Objetivo", level=1)
    doc.add_paragraph(
        "Formalizar las reglas que gobiernan el comportamiento operativo del sistema "
        "Tasf.B2B, asegurando trazabilidad con la situacion autentica del curso."
    )

    rules_groups = [
        ("2. Reglas de Simulacion", [
            ("RN-01", "La simulacion solo puede reconfigurarse cuando no esta en ejecucion (isRunning = false)."),
            ("RN-02", "El estado de simulacion debe ser consultable en todo momento via GET /api/simulation/state."),
            ("RN-03", "La velocidad de simulacion permite aceleracion controlada (1x a 10x)."),
            ("RN-04", "Existen 3 escenarios operativos: DAY_TO_DAY, PERIOD_SIMULATION, COLLAPSE_TEST."),
            ("RN-05", "La simulacion de periodo toma entre 30 y 90 minutos reales para simular 3, 5 o 7 dias."),
        ]),
        ("3. Reglas de Envios y Rutas", [
            ("RN-06", "Todo envio tiene origen, destino, fecha de registro, deadline y estado."),
            ("RN-07", "Un envio sin ruta planificada permanece en estado PENDING."),
            ("RN-08", "Un envio planificado pasa a IN_ROUTE al iniciar su primer tramo activo."),
            ("RN-09", "Un envio con deadline vencido y no entregado pasa a CRITICAL."),
            ("RN-10", "Al completar todos los TravelStops, el estado final es DELIVERED."),
            ("RN-11", "El plazo de entrega es de 1 dia maximo para mismo continente y 2 dias para distinto continente."),
        ]),
        ("4. Reglas de Capacidad y Colapso", [
            ("RN-12", "Un vuelo no puede superar su capacidad maxima (currentLoad <= maxCapacity)."),
            ("RN-13", "Estado de aeropuerto: NORMAL (< normalThreshold%), ALERTA (normal..warning), CRITICO (> warningThreshold%)."),
            ("RN-14", "Los umbrales de semaforo (verde/ambar/rojo) son parametros configurables."),
            ("RN-15", "Riesgo de colapso = 60% ocupacion promedio + 30% fraccion nodos CRITICO + 10% fraccion envios problematicos."),
            ("RN-16", "Colapso total se declara cuando > 90% de aeropuertos alcanzan CRITICO."),
            ("RN-17", "Capacidad de almacenamiento por aeropuerto: 500-800 maletas."),
            ("RN-18", "Capacidad de vuelo: 150-250 (intracontinental), 150-400 (intercontinental)."),
        ]),
        ("5. Reglas de Importacion", [
            ("RN-19", "La importacion masiva se ejecuta por job asincrono con seguimiento de estado."),
            ("RN-20", "Resultado de importacion satisface: filas_leidas = filas_importadas + filas_fallidas."),
        ]),
        ("6. Reglas de Reportes y Auditoria", [
            ("RN-21", "Reporte SLA se presenta por tipo de ruta, cliente (aerolinea) y nodo destino (ICAO)."),
            ("RN-22", "Exportacion de resultados soporta CSV y PDF."),
            ("RN-23", "Eventos de planificacion, replanificacion y entrega se registran en ShipmentAuditLog."),
        ]),
    ]

    for heading, rules in rules_groups:
        doc.add_heading(heading, level=1)
        add_table(doc, ["ID", "Regla"], [[rid, rdesc] for rid, rdesc in rules])

    return "15.ana.reglas.negocio.v01.docx", doc


# ============================================================================
#  16 - Trazabilidad LE Exigibles
# ============================================================================

def build_doc_16() -> tuple[str, Document]:
    doc = new_doc()
    add_cover(doc, "Documentacion de Analisis", "16. Trazabilidad de LE Exigibles (v01)")

    doc.add_heading("1. Objetivo", level=1)
    doc.add_paragraph(
        "Mapear las Listas de Exigencias tipo Exigible (LE-E) con las funcionalidades "
        "implementadas, indicando componentes backend y frontend."
    )

    doc.add_heading("2. Matriz de Trazabilidad", level=1)

    blocks = [
        ("2.1 Monitoreo (LE-001..LE-028)", [
            ["DashboardController", "Backend", "KPIs, estado del sistema, listado, busqueda, detalle nodo, red de rutas"],
            ["ShipmentController", "Backend", "Listado, detalle, envios overdue y critical"],
            ["AirportController", "Backend", "Listado y detalle de aeropuertos"],
            ["FlightController", "Backend", "Estado de vuelos"],
            ["OperationalAlertController", "Backend", "Gestion de alertas"],
            ["app/page.tsx (Dashboard)", "Frontend", "Mapa MapLibre, envios en ruta, KPIs, filtros, detalle envio/nodo"],
        ]),
        ("2.2 Gestion de Envios (LE-029..LE-044)", [
            ["ShipmentController", "Backend", "CRUD de envios, replanificacion"],
            ["ShipmentOrchestratorService", "Backend", "Orquestacion de creacion + planificacion"],
            ["RoutePlannerService", "Backend", "Planificacion de ruta con GA/ACO"],
            ["app/shipments/page.tsx", "Frontend", "Listado y gestion de envios"],
            ["app/import/page.tsx", "Frontend", "Importacion CSV y dataset masivo"],
        ]),
        ("2.3 Simulacion (LE-049..LE-057)", [
            ["SimulationController", "Backend", "Endpoints de config, control y resultados"],
            ["SimulationEngineService", "Backend", "Motor de ticks periodicos"],
            ["SimulationRuntimeService", "Backend", "Estado en memoria (pausa, velocidad, KPIs)"],
            ["DemandGenerationService", "Backend", "Generacion de volumen sintetico"],
            ["app/simulation/page.tsx", "Frontend", "Panel de configuracion, control, KPIs vivos"],
        ]),
        ("2.4 Reportes (LE-074)", [
            ["ReportsController", "Backend", "GET /api/reports/sla-compliance"],
            ["ReportingService", "Backend", "Calculo de SLA por dimensiones"],
            ["app/reports/page.tsx", "Frontend", "Visualizacion de SLA con barras"],
        ]),
    ]

    for heading, rows in blocks:
        doc.add_heading(heading, level=2)
        add_table(doc, ["Artefacto", "Capa", "Descripcion"], rows)

    doc.add_heading("3. Trazabilidad CU -> LE", level=1)
    add_table(doc,
        ["Caso de Uso", "LE"],
        [
            ["CU-01, CU-02 (Simulacion)", "LE-049..LE-057"],
            ["CU-03 (Monitoreo en mapa)", "LE-001..LE-028"],
            ["CU-04, CU-05, CU-06 (Envios)", "LE-029..LE-044"],
            ["CU-07, CU-08 (Reportes/Export)", "LE-074"],
        ],
    )

    return "16.ana.trazabilidad.lee.v01.docx", doc


# ============================================================================
#  21 - ISA Ampliado (Seleccion y Programacion de Algoritmos)
# ============================================================================

def build_doc_21() -> tuple[str, Document]:
    doc = new_doc()
    add_cover(doc, "ISA Ampliado", "Seleccion y Programacion de Algoritmos (v02)")

    doc.add_heading("1. Objetivo", level=1)
    doc.add_paragraph(
        "Documentar la seleccion, pseudocodigo, estructuras de datos y programacion "
        "completa de los dos algoritmos metaheuristicos del componente planificador, "
        "conforme a los requisitos no funcionales del curso: (a) dos soluciones algoritmicas "
        "en Lenguaje Java evaluadas por experimentacion numerica; (b) ambos del tipo "
        "metaheuristicos. Para equipo de 4 integrantes, se incluye tanto la metaheuristica "
        "de rutas como la asignacion de las rutas."
    )

    doc.add_heading("2. Algoritmos Seleccionados", level=1)
    add_table(doc,
        ["Aspecto", "Algoritmo 1: Genetic Algorithm (GA)", "Algoritmo 2: Ant Colony Optimization (ACO)"],
        [
            ["Tipo", "Metaheuristico evolutivo", "Metaheuristico de inteligencia de enjambre"],
            ["Archivo fuente", "algorithm/GeneticAlgorithm.java", "algorithm/AntColonyOptimization.java"],
            ["Lineas de codigo", "465", "478"],
            ["Interfaz comun", "RouteOptimizer", "RouteOptimizer"],
        ],
    )

    doc.add_heading("2.1 Justificacion de GA", level=2)
    bullet(doc, "Buen compromiso entre exploracion (diversidad genetica) y explotacion (seleccion de elite).")
    bullet(doc, "Adaptable a restricciones de capacidad de vuelo y ventanas temporales.")
    bullet(doc, "Codificacion natural: cromosoma = secuencia de TravelStops = ruta completa.")
    bullet(doc, "Robusto para espacios de solucion grandes con multiples aeropuertos y vuelos.")

    doc.add_heading("2.2 Justificacion de ACO", level=2)
    bullet(doc, "Feromonas refuerzan rutas prometedoras de forma incremental iteracion a iteracion.")
    bullet(doc, "Heuristica local basada en carga y ocupacion favorece rutas de baja saturacion (anticolapso).")
    bullet(doc, "Flexible ante cambios dinamicos del grafo (cancelaciones, replanificacion).")
    bullet(doc, "Complementa a GA explorando caminos alternativos via probabilidad de transicion.")

    # -- Interface --
    doc.add_heading("3. Interfaz Comun: RouteOptimizer", level=1)
    doc.add_paragraph(
        "Ambos algoritmos implementan la interfaz RouteOptimizer, que define el contrato "
        "comun para planificacion, replanificacion y evaluacion de rendimiento."
    )
    ga_interface = read_java_file(BACKEND_ALG / "RouteOptimizer.java")
    add_code(doc, ga_interface, font_size=7)

    # -- GA pseudocode --
    doc.add_heading("4. Algoritmo Genetico (GA) - Pseudocodigo", level=1)

    doc.add_heading("4.1 Estructura del cromosoma", level=2)
    doc.add_paragraph(
        "Un individuo (cromosoma) representa la ruta completa de un envio como una lista "
        "ordenada de TravelStop. Cada gen es una parada (aeropuerto + vuelo de llegada). "
        "Se usa el record Individual(List<TravelStop> stops, double fitness)."
    )

    doc.add_heading("4.2 Flujo principal", level=2)
    add_code(doc, """ALGORITMO GeneticAlgorithm.planRoute(shipment, availableFlights, airports)
ENTRADA:
  shipment         -- envio a planificar (origen, destino, maletas, deadline)
  availableFlights -- vuelos elegibles (SCHEDULED, con capacidad)
  airports         -- aeropuertos del sistema
SALIDA:
  bestRoute        -- lista ordenada de TravelStop

INICIO
  population <- initialize(shipment, availableFlights, airports)
  SI population esta vacia ENTONCES RETORNAR vacio

  PARA gen = 0 HASTA generations-1 HACER
      population <- evolve(population, shipment, availableFlights)
  FIN PARA

  RETORNAR individuo con maximo fitness en population
FIN""")

    doc.add_heading("4.3 Funcion de fitness", level=2)
    doc.add_paragraph("Base 1000 puntos con penalizaciones:")
    add_code(doc, """FUNCION fitness(stops, shipment)
  score <- 1000.0
  SI stops vacio RETORNAR 0.0

  transitHours <- SUMA de flight.transitTimeDays * 24 para cada stop con vuelo
  eta <- shipment.registrationDate + ceil(transitHours) horas

  SI eta > shipment.deadline ENTONCES
      score <- score - 500                    // penalizacion por incumplimiento deadline
  FIN SI

  PARA CADA stop EN stops HACER
      pct <- stop.airport.occupancyPct
      SI pct > warningThresholdPct ENTONCES
          score <- score - 2.5 * (pct - warningThresholdPct)  // aeropuerto saturado
      FIN SI
      SI stop.flight != null Y stop.flight.loadPct > 80 ENTONCES
          score <- score - 1.4 * (loadPct - 80)               // vuelo sobrecargado
      FIN SI
  FIN PARA

  score <- score - 20 * max(0, stops.size - 2)  // penalizacion por hops extras
  RETORNAR score
FIN""")

    doc.add_heading("4.4 Operadores geneticos", level=2)
    doc.add_paragraph("Evolucion (evolve): elitismo top 10% + seleccion aleatoria + crossover + mutacion.")
    add_code(doc, """FUNCION evolve(population, shipment, availableFlights)
  eliteCount <- ceil(population.size * 0.1)
  next <- population[0..eliteCount]     // conservar elite

  MIENTRAS |next| < populationSize HACER
      a <- seleccionar aleatorio de population
      b <- seleccionar aleatorio de population
      child <- crossover(a, b, shipment)
      SI random() < mutationRate ENTONCES
          child <- mutate(child, availableFlights, shipment)
      FIN SI
      next.agregar(Individual(child.stops, fitness(child.stops, shipment)))
  FIN MIENTRAS

  next.ordenar por fitness descendente
  RETORNAR next[0..populationSize]
FIN

CROSSOVER(parentA, parentB):
  Selecciona el padre con mayor fitness como base.

MUTACION(individual):
  Genera ruta alternativa via buildBestRoute y la sustituye si es factible.""")

    doc.add_heading("4.5 Parametros tuneados (perfil GA-P1)", level=2)
    add_table(doc,
        ["Parametro", "Rango explorado", "Valor optimo"],
        [
            ["populationSize", "20 - 100", "55"],
            ["generations", "10 - 50", "24"],
            ["mutationRate", "0.01 - 0.50", "0.05"],
            ["warningThresholdPct", "-", "90"],
        ],
    )

    # -- ACO pseudocode --
    doc.add_heading("5. Ant Colony Optimization (ACO) - Pseudocodigo", level=1)

    doc.add_heading("5.1 Estructura de feromonas", level=2)
    doc.add_paragraph(
        "Matriz de feromonas: Map<Long, Double> indexada por flightId. Cada entrada "
        "tau(f) representa la calidad acumulada del vuelo f como componente de rutas exitosas. "
        "Inicializada con initialPheromone = 1.0 para todos los vuelos."
    )

    doc.add_heading("5.2 Flujo principal", level=2)
    add_code(doc, """ALGORITMO AntColonyOptimization.planRoute(shipment, availableFlights, airports)
ENTRADA:
  shipment         -- envio a planificar
  availableFlights -- vuelos elegibles
  airports         -- aeropuertos del sistema
  config           -- numAnts, iterations, evaporationRate, alpha, beta
SALIDA:
  bestRoute        -- lista ordenada de TravelStop

INICIO
  initialize(availableFlights)   // tau(f) = 1.0 para todo f
  bestSolution <- vacio
  bestQuality <- -infinito

  PARA iter = 0 HASTA iterations-1 HACER
      antSolutions <- []

      // Cada hormiga construye una solucion
      PARA ant = 0 HASTA numAnts-1 HACER
          solution <- buildSolution(shipment, availableFlights, airports)
          SI solution no vacia ENTONCES
              antSolutions.agregar(solution)
              quality <- evaluateSolution(solution, shipment)
              SI quality > bestQuality ENTONCES
                  bestQuality <- quality
                  bestSolution <- solution
              FIN SI
          FIN SI
      FIN PARA

      // Actualizar feromonas
      evaporate()                           // tau(f) *= (1 - rho)
      updatePheromones(antSolutions, shipment)  // depositar en buenas rutas
  FIN PARA

  RETORNAR bestSolution
FIN""")

    doc.add_heading("5.3 Formula de probabilidad de transicion", level=2)
    doc.add_paragraph(
        "La probabilidad de seleccionar el vuelo f en la ruleta es:"
    )
    add_code(doc, """P(f) = [tau(f)]^alpha * [eta(f)]^beta / SUM_k [tau(k)]^alpha * [eta(k)]^beta

donde:
  tau(f) = feromona del vuelo f (experiencia acumulada)
  eta(f) = heuristica local = 1 / (1 + loadPct(f))
  alpha  = peso de feromona (importancia de experiencia colectiva)
  beta   = peso de heuristica (importancia de informacion local)""")

    doc.add_heading("5.4 Mecanismo de feromonas", level=2)
    add_code(doc, """EVAPORACION:
  PARA CADA arista f EN pheromones HACER
      tau(f) <- max(0.01, tau(f) * (1 - evaporationRate))
  FIN PARA

DEPOSITO (tras cada iteracion):
  Q <- 100.0 (constante de deposito)
  PARA CADA solution EN antSolutions HACER
      quality <- evaluateSolution(solution, shipment)
      SI quality <= 0 ENTONCES continuar  // solucion invalida
      delta <- Q / (1 + max(0, 1000 - quality))
      PARA CADA stop EN solution con vuelo HACER
          tau(stop.flight) <- tau(stop.flight) + delta
      FIN PARA
  FIN PARA""")

    doc.add_heading("5.5 Funcion de evaluacion de solucion", level=2)
    add_code(doc, """FUNCION evaluateSolution(stops, shipment)
  score <- 1000.0
  transitHours <- SUMA flight.transitTimeDays * 24

  SI eta > deadline ENTONCES score -= 420
  PARA CADA stop: SI occupancyPct > 90 ENTONCES score -= 2.2 * (pct - 90)
  PARA CADA stop con vuelo: SI loadPct > 80 ENTONCES score -= 1.1 * (loadPct - 80)
  score -= 14 * max(0, stops.size - 2)
  RETORNAR score
FIN""")

    doc.add_heading("5.6 Parametros tuneados", level=2)
    add_table(doc,
        ["Parametro", "Rango explorado", "Valor optimo"],
        [
            ["numAnts", "20 - 50", "20"],
            ["iterations", "30 - 100", "24"],
            ["evaporationRate (rho)", "0.02 - 0.40", "0.10"],
            ["alpha", "0.5 - 3.0", "1.1"],
            ["beta", "1.0 - 5.0", "2.1"],
            ["initialPheromone", "-", "1.0"],
        ],
    )

    # -- Structures diagram --
    doc.add_heading("6. Esquema de Estructuras", level=1)
    add_plantuml(doc, """@startuml
skinparam backgroundColor white
skinparam classAttributeIconSize 0

package "Componente Planificador" {

  interface RouteOptimizer {
    +planRoute(shipment, flights, airports): List<TravelStop>
    +replanRoute(shipment, failedStop, flights): List<TravelStop>
    +evaluatePerformance(shipments, from, to): OptimizationResult
  }

  class GeneticAlgorithm {
    -populationSize: int = 55
    -generations: int = 24
    -mutationRate: double = 0.05
    -warningThresholdPct: int = 90
    -initialize(): List<Individual>
    -evolve(): List<Individual>
    -crossover(): Individual
    -mutate(): Individual
    -fitness(): double
    -buildBestRoute(): List<TravelStop>
  }

  class AntColonyOptimization {
    -numAnts: int = 20
    -iterations: int = 24
    -evaporationRate: double = 0.10
    -alpha: double = 1.1
    -beta: double = 2.1
    -pheromones: Map<Long, Double>
    -initialize(): void
    -buildSolution(): List<TravelStop>
    -updatePheromones(): void
    -evaporate(): void
    -rouletteSelect(): Flight
    -evaluateSolution(): double
  }

  class RoutePlannerService {
    +planShipment(shipment, algorithm): List<TravelStop>
    +replanShipment(id): List<TravelStop>
    -resolveOptimizer(name): RouteOptimizer
    -buildDirectCandidate(): CandidatePlan
    -evaluateCandidate(): CandidatePlan
    -selectBestCandidate(): CandidatePlan
    -allocateFlightLoads(): void
    -eligibleFlightsForShipment(): List<Flight>
  }

  RouteOptimizer <|.. GeneticAlgorithm
  RouteOptimizer <|.. AntColonyOptimization
  RoutePlannerService --> RouteOptimizer : usa
}
@enduml""", "Figura 1. Estructura del Componente Planificador")

    # -- Route assignment --
    doc.add_heading("7. Asignacion de Rutas (equipo de 4)", level=1)
    doc.add_paragraph(
        "Ademas de las metaheuristicas de rutas, se implementa la asignacion de las rutas "
        "con reserva de capacidad de vuelo en RoutePlannerService:"
    )
    add_code(doc, """ESQUEMA DE ASIGNACION (RoutePlannerService)

1) Filtrar vuelos elegibles:
   - status = SCHEDULED
   - availableCapacity >= shipment.luggageCount
   - departure >= shipment.registrationDate

2) Generar candidatos:
   - Candidato DIRECTO: vuelo directo origen->destino (si existe)
   - Candidato MULTI-HOP: ruta optimizada por GA o ACO

3) Evaluar candidatos con funcion de costo:
   costo = etaHours
         + 0.75 * (stops.size - 2)           // hop penalty
         + SUM max(0, (loadPct-85)*0.08)      // load penalty por vuelo
         + 2.0 * count(occupancy >= warning%) // node penalty
         + [500 + 25*horasLate] si eta > deadline  // deadline penalty

4) Seleccionar mejor candidato (menor costo total)

5) Reservar capacidad por tramo (allocateFlightLoads):
   PARA CADA stop CON vuelo EN ruta HACER
       flight.currentLoad <- min(maxCapacity, currentLoad + luggageCount)
   FIN PARA

6) Persistir TravelStops y registrar evento en ShipmentAuditLog""")

    # -- Source code --
    doc.add_heading("8. Codigo Fuente: GeneticAlgorithm.java", level=1)
    doc.add_paragraph("A continuacion se presenta el codigo fuente completo del algoritmo genetico implementado en Java.")
    ga_code = read_java_file(BACKEND_ALG / "GeneticAlgorithm.java")
    add_code(doc, ga_code, font_size=7)

    doc.add_heading("9. Codigo Fuente: AntColonyOptimization.java", level=1)
    doc.add_paragraph("A continuacion se presenta el codigo fuente completo de Ant Colony Optimization implementado en Java.")
    aco_code = read_java_file(BACKEND_ALG / "AntColonyOptimization.java")
    add_code(doc, aco_code, font_size=7)

    # -- Pair programming --
    doc.add_heading("10. Programacion en Pares y Aportes", level=1)

    doc.add_heading("10.1 Par 1: A1 (Alonso Reyes Samaniego) + A2 (Rodrigo Holguin Huari)", level=2)
    add_table(doc,
        ["Integrante", "Aportes"],
        [
            ["A1 - Alonso Reyes Samaniego",
             "Integracion de flujos de simulacion con planificacion incremental. "
             "Validacion funcional de los 3 escenarios en el mapa interactivo. "
             "Coordinacion entre SimulationEngineService y RoutePlannerService."],
            ["A2 - Rodrigo Holguin Huari",
             "Ajustes de control runtime (pausa, velocidad, estado). "
             "Exposicion de resultados operativos y exportacion CSV/PDF. "
             "Evidencia de importacion/exportacion para pruebas de operacion."],
        ],
    )

    doc.add_heading("10.2 Par 2: A3 (Sebastian Badajoz Garcia Godos) + A4 (Adrian Picoy Cotrina)", level=2)
    add_table(doc,
        ["Integrante", "Aportes"],
        [
            ["A3 - Sebastian Badajoz Garcia Godos",
             "Especificacion y ajuste de componentes metaheuristicos (GA y ACO). "
             "Definicion de funcion de fitness, operadores geneticos y formula de feromonas. "
             "Criterio de evaluacion y comparacion tecnica interna (benchmark)."],
            ["A4 - Adrian Picoy Cotrina",
             "Implementacion de asignacion de capacidad por tramos (allocateFlightLoads). "
             "Persistencia de TravelStops y auditoria de planificacion. "
             "Soporte a robustez de importacion masiva y consistencia de datos."],
        ],
    )

    # -- Benchmark --
    doc.add_heading("11. Resultados de Experimentacion Numerica", level=1)
    doc.add_paragraph(
        "Se ejecutaron 204 corridas de benchmark sobre 9 escenarios oficiales. "
        "El ganador global fue Genetic Algorithm con el perfil GA-P1."
    )

    add_table(doc,
        ["Metrica", "Valor"],
        [
            ["Perfil ganador", "GA-P1"],
            ["Algoritmo ganador", "Genetic Algorithm"],
            ["Score promedio", "80.2649"],
            ["Collapse delay", "31.81 horas"],
            ["Completed", "100.00%"],
            ["Avg transit", "12.20 horas"],
            ["Corridas totales", "204"],
        ],
    )

    doc.add_heading("11.1 Analisis de sensibilidad de vuelos", level=2)
    add_table(doc,
        ["Variacion de vuelos", "Score", "Collapse delay (h)"],
        [
            ["-30%", "80.005", "33.31"],
            ["-15%", "79.710", "31.62"],
            ["Base (0%)", "80.455", "31.46"],
            ["+15%", "81.346", "32.59"],
            ["+30%", "80.527", "31.40"],
        ],
    )
    doc.add_paragraph(
        "El GA muestra robustez ante variaciones de frecuencia de vuelos, manteniendo "
        "score estable (~80) y delay de colapso consistente (~31-33 horas)."
    )

    return "21.dis.selec.algoritmos.v02.docx", doc


# ============================================================================
#  Main
# ============================================================================

def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    builders = [
        build_doc_01,
        build_doc_12,
        build_doc_13,
        build_doc_14,
        build_doc_15,
        build_doc_16,
        build_doc_21,
    ]

    print("Generando entregables DOCX...")
    for builder in builders:
        filename, doc = builder()
        path = OUT_DIR / filename
        doc.save(str(path))
        print(f"  OK: {path}")

    print(f"\n{len(builders)} documentos generados en {OUT_DIR}")


if __name__ == "__main__":
    main()
